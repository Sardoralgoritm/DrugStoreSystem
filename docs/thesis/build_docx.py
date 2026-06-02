"""
DrugstoreSystem BMI — Word hujjatini avtomatik yaratish
Shablon: BMI document change.docx asosida
Times New Roman 14pt, 1.5 interval, L=3cm R=1.5cm T=2cm B=2cm

Ishlatish:
    pip install python-docx Pillow
    python docs/thesis/build_docx.py
"""

import os, re
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image as PILImage

THESIS_DIR = os.path.dirname(os.path.abspath(__file__))
IMAGES_DIR = os.path.join(THESIS_DIR, "images")
OUTPUT     = os.path.join(THESIS_DIR, "DrugstoreSystem-BMI.docx")
TEMPLATE   = r"C:\Users\sardo\Downloads\Telegram Desktop\BMI document change.docx"

MD_FILES = [
    "00-titul-annotatsiya.md",
    "01-kirish.md",
    "02-bob1-01-dorixona-tizimlari.md",
    "02-bob1-02-qidiruv-haversine.md",
    "02-bob1-03-mavjud-platformalar.md",
    "03-bob2-01-arxitektura.md",
    "03-bob2-02-database.md",
    "04-bob3-01-imkoniyatlar.md",
    "04-bob3-02-algoritmlar.md",
    "05-bob4-01-ergonomika.md",
    "05-bob4-02-yongin.md",
    "06-xulosa.md",
    "07-adabiyotlar.md",
    "08-ilovalar.md",
]

# ── helpers ───────────────────────────────────────────────────────────────────

def _rfonts(run, name="Times New Roman"):
    rPr = run._r.get_or_add_rPr()
    rf = OxmlElement("w:rFonts")
    rf.set(qn("w:ascii"), name)
    rf.set(qn("w:hAnsi"), name)
    rf.set(qn("w:cs"),    name)
    rPr.insert(0, rf)

def fmt_run(run, size=14, bold=False, italic=False, mono=False):
    name = "Courier New" if mono else "Times New Roman"
    run.font.name   = name
    run.font.size   = Pt(size)
    run.font.bold   = bold
    run.font.italic = italic
    _rfonts(run, name)

def fmt_para(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
             sb=0, sa=6, ls=1.5, indent=1.25):
    f = p.paragraph_format
    f.alignment          = align
    f.space_before       = Pt(sb)
    f.space_after        = Pt(sa)
    f.line_spacing_rule  = WD_LINE_SPACING.MULTIPLE
    f.line_spacing       = ls
    f.first_line_indent  = Cm(indent) if indent else None

def new_para(doc, text="", align=WD_ALIGN_PARAGRAPH.JUSTIFY,
             sb=0, sa=6, ls=1.5, indent=1.25,
             size=14, bold=False, italic=False, mono=False):
    p = doc.add_paragraph()
    fmt_para(p, align=align, sb=sb, sa=sa, ls=ls, indent=indent)
    if text:
        r = p.add_run(text)
        fmt_run(r, size=size, bold=bold, italic=italic, mono=mono)
    return p

# ── Uzbek apostrophe + quotation fix ─────────────────────────────────────────

def fix_uzb(text: str) -> str:
    """Replace straight apostrophes and quotes with correct Unicode variants."""
    text = text.replace("o'", "o‘").replace("O'", "O‘")
    text = text.replace("g'", "g‘").replace("G'", "G‘")
    text = text.replace("'", "’")
    text = re.sub(r'"([^"]+)"', "“\\1”", text)
    return text

# ── inline markdown: **bold**, *italic*, `code` ───────────────────────────────

INLINE = re.compile(r"\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`")

def add_inline(p, text, base_size=14):
    text = fix_uzb(text)
    last = 0
    for m in INLINE.finditer(text):
        if m.start() > last:
            r = p.add_run(text[last:m.start()])
            fmt_run(r, size=base_size)
        if m.group(1):
            r = p.add_run(m.group(1))
            fmt_run(r, size=base_size, bold=True)
        elif m.group(2):
            r = p.add_run(m.group(2))
            fmt_run(r, size=base_size, italic=True)
        elif m.group(3):
            r = p.add_run(m.group(3))
            fmt_run(r, size=12, mono=True)
        last = m.end()
    if last < len(text):
        r = p.add_run(text[last:])
        fmt_run(r, size=base_size)

# ── page numbers ──────────────────────────────────────────────────────────────

def add_page_numbers(doc):
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.clear()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        fmt_run(run, size=12)
        for tag, txt in [("begin", None), (None, "PAGE"), ("end", None)]:
            if tag:
                el = OxmlElement("w:fldChar")
                el.set(qn("w:fldCharType"), tag)
                run._r.append(el)
            else:
                el = OxmlElement("w:instrText")
                el.text = txt
                run._r.append(el)

# ── title page ────────────────────────────────────────────────────────────────

def add_title_page(doc):
    def c(text, size=14, bold=False, sa=6):
        p = new_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER,
                     sb=0, sa=sa, ls=1.5, indent=0,
                     size=size, bold=bold)
        r = p.add_run(text)
        fmt_run(r, size=size, bold=bold)
        return p

    c("O'ZBEKISTON RESPUBLIKASI OLIY VA O'RTA MAXSUS TA'LIM VAZIRLIGI", 12, bold=True, sa=4)
    c("ALFRAGANUS UNIVERSITETI", 14, bold=True, sa=4)
    c('"Raqamli texnologiyalari" kafedrasi', 13, sa=36)
    c("BITIRUV MALAKAVIY ISHI", 16, bold=True, sa=8)
    c("mavzu:", 14, sa=4)
    c("DORIXONALARDAN DORI BUYURTMA BERISHNING ENG YAQIN", 15, bold=True, sa=4)
    c("YECHIMINI TOPIB BERUVCHI DASTURIY VOSITASINI ISHLAB CHIQISH", 15, bold=True, sa=36)

    for label, value in [
        ("Bajardi:",        "Sharipova Moxichexra Oltinovna"),
        ("Guruh:",          "M313-22 Dlo'"),
        ("Ilmiy rahbar:",   "Hamroev Alisher Shodmonqulovich, dotsent"),
        ("Kafedra mudiri:", "________________"),
    ]:
        p = new_para(doc, align=WD_ALIGN_PARAGRAPH.RIGHT, sb=0, sa=4, ls=1.5, indent=0)
        r1 = p.add_run(label + "  ")
        fmt_run(r1, size=13, bold=True)
        r2 = p.add_run(value)
        fmt_run(r2, size=13)

    p = new_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, sb=48, sa=0, ls=1.5, indent=0)
    r = p.add_run("Toshkent — 2025")
    fmt_run(r, size=14)
    doc.add_page_break()

# ── TOC ───────────────────────────────────────────────────────────────────────

def add_toc(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(14)
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run("MUNDARIJA")
    fmt_run(r, size=14, bold=True)

    toc_entries = [
        ("ANNOTATSIYA", "3", True),
        ("KIRISH", "6", True),
        ("I BOB. DORIXONA AXBOROT TIZIMLARINI RAQAMLASHTIRISH VA OPTIMALLASHTIRISH ALGORITMLARINING NAZARIY ASOSLARI", "12", True),
        ("    1.1. Zamonaviy dorixona boshqaruv tizimlari va farmatsevtika sohasini raqamlashtirish", "12", False),
        ("    1.2. Qidiruv algoritmlari va geodezik masofani hisoblash nazariy asoslari", "17", False),
        ("    1.3. Mavjud dorixona platformalari va qidiruv tizimlarining qiyosiy tahlili", "22", False),
        ("II BOB. DORIXONA QIDIRUV TIZIMINI LOYIHALASH VA ISHLAB CHIQISH", "28", True),
        ("    2.1. Tizim arxitekturasi va zaruriy texnologiyalarni tanlash", "28", False),
        ("    2.2. Ma'lumotlar bazasini loyihalash va integratsiyalash", "35", False),
        ("III BOB. LOYIHANING AMALIY TATBIQI VA ALGORITMLARINING SAMARADORLIGI", "43", True),
        ("    3.1. Tizim imkoniyatlari bilan tanishish", "43", False),
        ("    3.2. Qidiruv va optimallashtirish algoritmlarining samaradorligi", "50", False),
        ("IV BOB. HAYOT FAOLIYATI XAVFSIZLIGI", "55", True),
        ("    4.1. Kompyuter bilan ishlashda ergonomika va sog'liqni saqlash", "55", False),
        ("    4.2. Yong'in xavfsizligi", "60", False),
        ("XULOSA", "65", True),
        ("FOYDALANILGAN ADABIYOTLAR", "67", True),
        ("DASTUR ILOVASI", "70", True),
    ]

    for title, page_num, bold in toc_entries:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(3)
        p.paragraph_format.first_line_indent = None
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = 1.5

        pPr = p._p.get_or_add_pPr()
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "right")
        tab.set(qn("w:leader"), "dot")
        tab.set(qn("w:pos"), "9355")
        tabs.append(tab)
        pPr.append(tabs)

        r1 = p.add_run(title)
        fmt_run(r1, size=13, bold=bold)
        r2 = p.add_run("\t" + page_num)
        fmt_run(r2, size=13, bold=bold)

    doc.add_page_break()

# ── image ─────────────────────────────────────────────────────────────────────

def add_image(doc, img_path, caption=None):
    if not os.path.exists(img_path):
        new_para(doc, f"[Rasm topilmadi: {os.path.basename(img_path)}]",
                 align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, size=11, indent=0)
        return
    try:
        with PILImage.open(img_path) as im:
            w_px, h_px = im.size
            dpi = im.info.get("dpi", (96, 96))
            dpi_x = dpi[0] if isinstance(dpi, (tuple, list)) else 96
            w_cm = w_px / dpi_x * 2.54
            if w_cm > 14.5:
                w_cm = 14.5
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after  = Pt(2)
        p.paragraph_format.first_line_indent = None
        p.add_run().add_picture(img_path, width=Cm(w_cm))
    except Exception as e:
        new_para(doc, f"[Rasm xatosi: {e}]",
                 align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, size=11, indent=0)
        return
    if caption:
        cp = new_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, sb=0, sa=10,
                      ls=1.5, indent=0, size=12, italic=True)
        r = cp.add_run(caption)
        fmt_run(r, size=12, italic=True)

# ── code block ────────────────────────────────────────────────────────────────

def add_code(doc, code):
    for line in code.splitlines():
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing  = 1.15
        p.paragraph_format.first_line_indent = None
        p.paragraph_format.left_indent   = Cm(1.0)
        r = p.add_run(line or " ")
        fmt_run(r, size=10, mono=True)

# ── table ─────────────────────────────────────────────────────────────────────

def add_table(doc, lines):
    rows = []
    for line in lines:
        if re.match(r"^\s*\|[-: |]+\|\s*$", line):
            continue
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        rows.append(cells)
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    tbl = doc.add_table(rows=len(rows), cols=ncols)
    tbl.style = "Table Grid"
    for ri, row_data in enumerate(rows):
        row = tbl.rows[ri]
        for ci in range(ncols):
            cell_text = row_data[ci] if ci < len(row_data) else ""
            cell_text = cell_text.replace("**", "")
            cell = row.cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            p.paragraph_format.line_spacing  = 1.15
            p.paragraph_format.first_line_indent = None
            r = p.add_run(cell_text)
            fmt_run(r, size=12, bold=(ri == 0))
    doc.add_paragraph()

# ── heading ───────────────────────────────────────────────────────────────────

def add_heading(doc, text, level):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.5
    text = fix_uzb(text)
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after  = Pt(8)
        r = p.add_run(text.upper())
    elif level == 2:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after  = Pt(6)
        r = p.add_run(text)
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after  = Pt(4)
        r = p.add_run(text)
    fmt_run(r, size=14, bold=True)

# ── markdown parser ───────────────────────────────────────────────────────────

def process_md(doc, filepath):
    with open(filepath, encoding="utf-8-sig") as f:
        lines = f.read().splitlines()

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip().lstrip("﻿")

        if stripped.startswith("### "):
            add_heading(doc, stripped[4:], 3); i += 1; continue
        if stripped.startswith("## "):
            add_heading(doc, stripped[3:], 2); i += 1; continue
        if stripped.startswith("# "):
            add_heading(doc, stripped[2:], 1); i += 1; continue

        m = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", stripped)
        if m:
            caption  = m.group(1)
            rel_path = m.group(2).replace("/", os.sep)
            add_image(doc, os.path.join(THESIS_DIR, rel_path), caption)
            i += 1; continue

        if re.match(r"^\*\*\d+\.\d+", stripped):
            p = new_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER,
                         sb=0, sa=10, ls=1.5, indent=0)
            add_inline(p, stripped, base_size=12)
            i += 1; continue

        if stripped.startswith("```"):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i]); i += 1
            i += 1
            add_code(doc, "\n".join(code_lines))
            doc.add_paragraph(); continue

        if stripped in ("***", "---", "* * *", "- - -"):
            i += 1; continue

        if stripped.startswith("|"):
            tbl_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                tbl_lines.append(lines[i]); i += 1
            add_table(doc, tbl_lines); continue

        m = re.match(r"^(\d+)\.\s+(.*)", stripped)
        if m:
            p = doc.add_paragraph()
            fmt_para(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                     sb=0, sa=3, ls=1.5, indent=0)
            p.paragraph_format.left_indent = Cm(1.25)
            add_inline(p, m.group(1) + ". " + m.group(2))
            i += 1; continue

        if re.match(r"^[-*]\s+", stripped):
            text = re.sub(r"^[-*]\s+", "", stripped)
            p = doc.add_paragraph()
            fmt_para(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                     sb=0, sa=3, ls=1.5, indent=0)
            p.paragraph_format.left_indent = Cm(1.25)
            add_inline(p, "- " + text)
            i += 1; continue

        if not stripped:
            i += 1; continue

        para_lines = [stripped]
        i += 1
        while i < len(lines):
            nxt = lines[i].strip().lstrip("﻿")
            if (not nxt or
                nxt.startswith("#") or
                nxt.startswith("|") or
                nxt.startswith("```") or
                nxt.startswith("![") or
                nxt in ("***", "---") or
                re.match(r"^(\d+)\.\s", nxt) or
                re.match(r"^[-*]\s", nxt)):
                break
            para_lines.append(nxt)
            i += 1

        full = " ".join(para_lines)
        p = doc.add_paragraph()
        fmt_para(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                 sb=0, sa=6, ls=1.5, indent=1.25)
        add_inline(p, full)

# ── main ─────────────────────────────────────────────────────────────────────

def main():
    doc = Document()

    for s in doc.sections:
        s.page_width    = Cm(21)
        s.page_height   = Cm(29.7)
        s.left_margin   = Cm(3.0)
        s.right_margin  = Cm(1.5)
        s.top_margin    = Cm(2.0)
        s.bottom_margin = Cm(2.0)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(14)

    add_title_page(doc)
    add_page_numbers(doc)
    add_toc(doc)

    for md in MD_FILES:
        path = os.path.join(THESIS_DIR, md)
        if not os.path.exists(path):
            print(f"SKIP: {md}"); continue
        print(f"OK:   {md}")
        process_md(doc, path)
        doc.add_page_break()

    doc.save(OUTPUT)
    mb = os.path.getsize(OUTPUT) / 1024 / 1024
    print(f"\nSaved: {OUTPUT}  ({mb:.1f} MB)")

if __name__ == "__main__":
    main()
